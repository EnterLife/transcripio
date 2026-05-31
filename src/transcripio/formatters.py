from __future__ import annotations

from io import BytesIO

from transcripio.models import TranscriptSegment, TranscriptionResult
from transcripio.storage import result_to_dict

import json


def to_txt(segments: list[TranscriptSegment]) -> str:
    lines: list[str] = []
    for segment in segments:
        prefix = f"{segment.speaker}: " if segment.speaker else ""
        lines.append(f"{_format_time(segment.start)} {_format_time(segment.end)}  {prefix}{segment.text}")
    return "\n".join(lines)


def to_srt(segments: list[TranscriptSegment]) -> str:
    blocks: list[str] = []
    for index, segment in enumerate(segments, start=1):
        speaker = f"{segment.speaker}: " if segment.speaker else ""
        blocks.append(
            "\n".join(
                [
                    str(index),
                    f"{_format_srt_time(segment.start)} --> {_format_srt_time(segment.end)}",
                    f"{speaker}{segment.text}",
                ]
            )
        )
    return "\n\n".join(blocks)


def to_vtt(segments: list[TranscriptSegment]) -> str:
    blocks = ["WEBVTT", ""]
    for segment in segments:
        speaker = f"{segment.speaker}: " if segment.speaker else ""
        blocks.append(
            "\n".join(
                [
                    f"{_format_vtt_time(segment.start)} --> {_format_vtt_time(segment.end)}",
                    f"{speaker}{segment.text}",
                ]
            )
        )
        blocks.append("")
    return "\n".join(blocks).strip() + "\n"


def to_docx(result: TranscriptionResult) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Run: pip install -r requirements.txt") from exc

    document = Document()
    document.add_heading("Transcript", level=1)
    document.add_paragraph(f"Source: {result.source_name}")
    if result.language:
        document.add_paragraph(f"Language: {result.language}")
    if result.duration is not None:
        document.add_paragraph(f"Duration: {_format_duration(result.duration)}")

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Start"
    headers[1].text = "End"
    headers[2].text = "Speaker"
    headers[3].text = "Text"

    for segment in result.segments:
        cells = table.add_row().cells
        cells[0].text = _format_plain_time(segment.start)
        cells[1].text = _format_plain_time(segment.end)
        cells[2].text = segment.speaker or ""
        cells[3].text = segment.text

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def to_json(result: TranscriptionResult) -> str:
    return json.dumps(result_to_dict(result), ensure_ascii=False, indent=2)


def _format_time(seconds: float) -> str:
    hours, minutes, sec, _millis = _split_timestamp(seconds)
    return f"[{hours:02d}:{minutes:02d}:{sec:02d}]"


def _format_srt_time(seconds: float) -> str:
    hours, minutes, sec, millis = _split_timestamp(seconds)
    return f"{hours:02d}:{minutes:02d}:{sec:02d},{millis:03d}"


def _format_vtt_time(seconds: float) -> str:
    hours, minutes, sec, millis = _split_timestamp(seconds)
    return f"{hours:02d}:{minutes:02d}:{sec:02d}.{millis:03d}"


def _format_plain_time(seconds: float) -> str:
    hours, minutes, sec, _millis = _split_timestamp(seconds)
    return f"{hours:02d}:{minutes:02d}:{sec:02d}"


def _format_duration(seconds: float) -> str:
    hours, minutes, sec, _millis = _split_timestamp(seconds)
    return f"{hours:02d}:{minutes:02d}:{sec:02d}"


def _split_timestamp(seconds: float) -> tuple[int, int, int, int]:
    total_millis = int(round(seconds * 1000))
    whole, millis = divmod(total_millis, 1000)
    minutes, sec = divmod(whole, 60)
    hours, minutes = divmod(minutes, 60)
    return hours, minutes, sec, millis
